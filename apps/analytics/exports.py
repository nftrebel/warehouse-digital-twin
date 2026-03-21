"""
Сервис экспорта аналитических отчётов (DOCX и Excel).

Содержание отчёта:
  1. Ключевые показатели (KPI)
  2. Распределение партий и заказов по этапам
  3. События по типам и источникам
  4. Длительность этапов
  5. Детализация отгруженных заказов за период
  6. Детализация просроченных заказов
"""

import io
from datetime import datetime, timedelta
from decimal import Decimal

from django.db.models import Count
from django.utils import timezone

from apps.inventory.models import Batch
from apps.orders.models import CustomerOrder
from apps.events.models import ProcessEvent


# ======================================================================
# Сбор данных
# ======================================================================

def _get_report_data(date_from=None, date_to=None):
    events_qs = ProcessEvent.objects.filter(processing_status='applied')
    if date_from:
        events_qs = events_qs.filter(event_time__gte=date_from)
    if date_to:
        events_qs = events_qs.filter(event_time__lte=date_to)

    now = timezone.now()

    total_events = events_qs.count()
    events_by_type = list(
        events_qs.values('event_type_code')
        .annotate(count=Count('event_id')).order_by('-count')
    )
    events_by_source = list(
        events_qs.values('source_system')
        .annotate(count=Count('event_id')).order_by('-count')
    )

    # Shipped orders in period
    shipped_order_ids = list(
        events_qs.filter(event_type_code='shipment.dispatched')
        .exclude(order_id=None)
        .values_list('order_id', flat=True).distinct()
    )
    shipped_orders = list(
        CustomerOrder.objects.filter(order_id__in=shipped_order_ids)
        .prefetch_related('lines__product')
    )

    # Overdue orders
    overdue_orders = list(
        CustomerOrder.objects.filter(planned_ship_date__lt=now)
        .exclude(current_stage_code='shipped')
        .prefetch_related('lines__product')
    )

    batch_stages = list(
        Batch.objects.values('current_stage_code')
        .annotate(count=Count('batch_id')).order_by('current_stage_code')
    )
    order_stages = list(
        CustomerOrder.objects.values('current_stage_code')
        .annotate(count=Count('order_id')).order_by('current_stage_code')
    )

    def calc_avg(from_type, to_type, key_field):
        qs_from = dict(
            events_qs.filter(event_type_code=from_type)
            .exclude(**{key_field: None}).values_list(key_field, 'event_time')
        )
        qs_to = dict(
            events_qs.filter(event_type_code=to_type)
            .exclude(**{key_field: None}).values_list(key_field, 'event_time')
        )
        durations = []
        for k, t0 in qs_from.items():
            if k in qs_to:
                d = (qs_to[k] - t0).total_seconds()
                if d > 0:
                    durations.append(d)
        if durations:
            return _fmt(sum(durations) / len(durations)), len(durations)
        return '—', 0

    batch_transitions = [
        ('batch.received', 'batch.placed', 'Приёмка → Размещение', 'batch_id'),
        ('batch.received', 'batch.reserved', 'Приёмка → Резервирование', 'batch_id'),
        ('batch.placed', 'batch.reserved', 'Размещение → Резервирование', 'batch_id'),
    ]
    order_transitions = [
        ('order.created', 'order.picking_started', 'Создан → Комплектация', 'order_id'),
        ('order.picking_started', 'order.assembled', 'Комплектация → Собран', 'order_id'),
        ('order.assembled', 'shipment.dispatched', 'Собран → Отгрузка', 'order_id'),
        ('order.created', 'shipment.dispatched', 'Создан → Отгрузка (полный цикл)', 'order_id'),
    ]

    batch_dur = [{'label': l, 'avg': calc_avg(f, t, k)[0], 'count': calc_avg(f, t, k)[1]}
                 for f, t, l, k in batch_transitions if calc_avg(f, t, k)[1] > 0]
    order_dur = [{'label': l, 'avg': calc_avg(f, t, k)[0], 'count': calc_avg(f, t, k)[1]}
                 for f, t, l, k in order_transitions if calc_avg(f, t, k)[1] > 0]

    avg_rp, _ = calc_avg('batch.received', 'batch.placed', 'batch_id')
    avg_ps, _ = calc_avg('order.picking_started', 'shipment.dispatched', 'order_id')

    return {
        'total_events': total_events,
        'events_by_type': events_by_type,
        'events_by_source': events_by_source,
        'shipped_orders': shipped_orders,
        'overdue_orders': overdue_orders,
        'batch_stages': batch_stages,
        'order_stages': order_stages,
        'batch_dur': batch_dur,
        'order_dur': order_dur,
        'avg_rp': avg_rp,
        'avg_ps': avg_ps,
        'date_from': date_from,
        'date_to': date_to,
        'generated_at': timezone.now(),
    }


def _fmt(sec):
    if sec < 60: return f'{int(sec)} сек'
    if sec < 3600: return f'{int(sec // 60)} мин'
    if sec < 86400: return f'{int(sec // 3600)} ч {int((sec % 3600) // 60)} мин'
    return f'{int(sec // 86400)} дн {int((sec % 86400) // 3600)} ч'


def _period(data):
    df, dt = data['date_from'], data['date_to']
    if df and dt: return f"{df.strftime('%d.%m.%Y')} — {dt.strftime('%d.%m.%Y')}"
    if df: return f"с {df.strftime('%d.%m.%Y')}"
    if dt: return f"по {dt.strftime('%d.%m.%Y')}"
    return "За всё время"


def _delay_str(now, planned):
    if not planned:
        return '—'
    delta = now - planned
    return f'{delta.days} дн {delta.seconds // 3600} ч'


def _qty_str(val):
    """Format Decimal removing trailing zeros."""
    result = f'{val:f}'
    if '.' in result:
        result = result.rstrip('0').rstrip('.')
    return result


# ======================================================================
# DOCX
# ======================================================================

def generate_docx(date_from=None, date_to=None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = _get_report_data(date_from, date_to)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Title
    _docx_centered(doc, 'Аналитический отчёт', 18, True, (0x1a, 0x1f, 0x36))
    _docx_centered(doc, 'Цифровой двойник складской логистики', 12, False, (0x7b, 0x80, 0x94))
    _docx_centered(doc, f'Период: {_period(data)}', 10, False, (0x3c, 0x40, 0x43))
    _docx_centered(doc, f'Сформирован: {data["generated_at"].strftime("%d.%m.%Y %H:%M")}', 9, False, (0x99, 0x99, 0x99))
    doc.add_paragraph()

    # KPI
    doc.add_heading('Ключевые показатели', level=1)
    _docx_table(doc, [
        ['Показатель', 'Значение'],
        ['Всего событий за период', str(data['total_events'])],
        ['Ср. время приёмка → размещение', data['avg_rp']],
        ['Ср. время комплектация → отгрузка', data['avg_ps']],
        ['Отгружено заказов за период', str(len(data['shipped_orders']))],
        ['Просроченных заказов', str(len(data['overdue_orders']))],
    ])

    # Batch stages
    doc.add_heading('Распределение партий по этапам', level=1)
    _docx_table(doc, [['Этап', 'Количество']] + [
        [s['current_stage_code'], str(s['count'])] for s in data['batch_stages']
    ])

    # Order stages
    doc.add_heading('Распределение заказов по этапам', level=1)
    _docx_table(doc, [['Этап', 'Количество']] + [
        [s['current_stage_code'], str(s['count'])] for s in data['order_stages']
    ])

    # Events by type
    doc.add_heading('События по типам', level=1)
    _docx_table(doc, [['Тип события', 'Количество']] + [
        [e['event_type_code'], str(e['count'])] for e in data['events_by_type']
    ])

    # Events by source
    doc.add_heading('События по источникам', level=1)
    _docx_table(doc, [['Источник', 'Количество']] + [
        [e['source_system'], str(e['count'])] for e in data['events_by_source']
    ])

    # Durations
    if data['batch_dur']:
        doc.add_heading('Длительность этапов партий', level=1)
        _docx_table(doc, [['Переход', 'Среднее время', 'Кол-во']] + [
            [d['label'], d['avg'], str(d['count'])] for d in data['batch_dur']
        ])
    if data['order_dur']:
        doc.add_heading('Длительность этапов заказов', level=1)
        _docx_table(doc, [['Переход', 'Среднее время', 'Кол-во']] + [
            [d['label'], d['avg'], str(d['count'])] for d in data['order_dur']
        ])

    # Shipped orders detail
    doc.add_page_break()
    doc.add_heading('Детализация отгруженных заказов', level=1)
    if data['shipped_orders']:
        p = doc.add_paragraph()
        run = p.add_run(f'Отгружено заказов за период: {len(data["shipped_orders"])}')
        run.bold = True
        run.font.color.rgb = RGBColor(0x13, 0x73, 0x33)

        _docx_table(doc, [['Номер', 'Приоритет', 'Позиций', 'Создан']] + [
            [
                o.order_number,
                o.get_priority_code_display(),
                str(o.lines.count()),
                o.created_at.strftime('%d.%m.%Y %H:%M'),
            ] for o in data['shipped_orders']
        ])
        doc.add_paragraph()

        for order in data['shipped_orders']:
            doc.add_heading(f'Заказ {order.order_number}', level=2)
            lines = order.lines.select_related('product').all()
            if lines:
                _docx_table(doc, [['Товар', 'SKU', 'Количество']] + [
                    [l.product.product_name, l.product.sku_code, _qty_str(l.requested_qty)]
                    for l in lines
                ])
            doc.add_paragraph()
    else:
        doc.add_paragraph('За выбранный период отгруженных заказов нет.')

    # Overdue orders detail
    doc.add_page_break()
    doc.add_heading('Детализация просроченных заказов', level=1)
    now = data['generated_at']
    if data['overdue_orders']:
        p = doc.add_paragraph()
        run = p.add_run(f'Просроченных заказов: {len(data["overdue_orders"])}')
        run.bold = True
        run.font.color.rgb = RGBColor(0xc5, 0x22, 0x1f)

        _docx_table(doc, [['Номер', 'Приоритет', 'Этап', 'Плановая отгрузка', 'Просрочка']] + [
            [
                o.order_number,
                o.get_priority_code_display(),
                o.get_current_stage_code_display(),
                o.planned_ship_date.strftime('%d.%m.%Y %H:%M') if o.planned_ship_date else '—',
                _delay_str(now, o.planned_ship_date),
            ] for o in data['overdue_orders']
        ])
        doc.add_paragraph()

        for order in data['overdue_orders']:
            doc.add_heading(f'Заказ {order.order_number}', level=2)
            lines = order.lines.select_related('product').all()
            if lines:
                _docx_table(doc, [['Товар', 'SKU', 'Количество']] + [
                    [l.product.product_name, l.product.sku_code, _qty_str(l.requested_qty)]
                    for l in lines
                ])
            doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run('Просроченных заказов не обнаружено.')
        run.font.color.rgb = RGBColor(0x13, 0x73, 0x33)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_centered(doc, text, size, bold, rgb):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*rgb)


def _docx_table(doc, rows):
    if len(rows) < 2:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Light Grid Accent 1')
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = str(cell)
    doc.add_paragraph()


# ======================================================================
# EXCEL
# ======================================================================

def generate_xlsx(date_from=None, date_to=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    data = _get_report_data(date_from, date_to)
    wb = Workbook()

    # Styles
    header_font = Font(name='Arial', size=10, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1A1F36', end_color='1A1F36', fill_type='solid')
    cell_font = Font(name='Arial', size=10)
    title_font = Font(name='Arial', size=14, bold=True, color='1A1F36')
    subtitle_font = Font(name='Arial', size=10, color='7B8094')
    green_font = Font(name='Arial', size=10, bold=True, color='137333')
    red_font = Font(name='Arial', size=10, bold=True, color='C5221F')
    thin_border = Border(
        left=Side(style='thin', color='E3E6EC'),
        right=Side(style='thin', color='E3E6EC'),
        top=Side(style='thin', color='E3E6EC'),
        bottom=Side(style='thin', color='E3E6EC'),
    )
    alt_fill = PatternFill(start_color='F8F9FB', end_color='F8F9FB', fill_type='solid')

    def write_header(ws, row, headers):
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = Alignment(horizontal='left', vertical='center')
            c.border = thin_border

    def write_row(ws, row, values, is_alt=False):
        for col, v in enumerate(values, 1):
            c = ws.cell(row=row, column=col, value=v)
            c.font = cell_font
            c.border = thin_border
            if is_alt:
                c.fill = alt_fill

    def auto_width(ws):
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_len + 4, 45)

    period = _period(data)
    now = data['generated_at']

    # ---- Sheet 1: KPI ----
    ws = wb.active
    ws.title = 'KPI'
    ws.cell(row=1, column=1, value='Аналитический отчёт — Цифровой двойник').font = title_font
    ws.cell(row=2, column=1, value=f'Период: {period}').font = subtitle_font
    ws.cell(row=3, column=1, value=f'Сформирован: {now.strftime("%d.%m.%Y %H:%M")}').font = subtitle_font

    write_header(ws, 5, ['Показатель', 'Значение'])
    kpi = [
        ('Всего событий за период', data['total_events']),
        ('Ср. время приёмка → размещение', data['avg_rp']),
        ('Ср. время комплектация → отгрузка', data['avg_ps']),
        ('Отгружено заказов за период', len(data['shipped_orders'])),
        ('Просроченных заказов', len(data['overdue_orders'])),
    ]
    for i, (label, val) in enumerate(kpi):
        write_row(ws, 6 + i, [label, val], i % 2 == 1)
    auto_width(ws)

    # ---- Sheet 2: Партии по этапам ----
    ws2 = wb.create_sheet('Партии по этапам')
    write_header(ws2, 1, ['Этап', 'Количество'])
    for i, s in enumerate(data['batch_stages']):
        write_row(ws2, 2 + i, [s['current_stage_code'], s['count']], i % 2 == 1)
    auto_width(ws2)

    # ---- Sheet 3: Заказы по этапам ----
    ws3 = wb.create_sheet('Заказы по этапам')
    write_header(ws3, 1, ['Этап', 'Количество'])
    for i, s in enumerate(data['order_stages']):
        write_row(ws3, 2 + i, [s['current_stage_code'], s['count']], i % 2 == 1)
    auto_width(ws3)

    # ---- Sheet 4: События по типам ----
    ws4 = wb.create_sheet('События по типам')
    write_header(ws4, 1, ['Тип события', 'Количество'])
    for i, e in enumerate(data['events_by_type']):
        write_row(ws4, 2 + i, [e['event_type_code'], e['count']], i % 2 == 1)
    auto_width(ws4)

    # ---- Sheet 5: События по источникам ----
    ws5 = wb.create_sheet('События по источникам')
    write_header(ws5, 1, ['Источник', 'Количество'])
    for i, e in enumerate(data['events_by_source']):
        write_row(ws5, 2 + i, [e['source_system'], e['count']], i % 2 == 1)
    auto_width(ws5)

    # ---- Sheet 6: Длительности ----
    ws6 = wb.create_sheet('Длительность этапов')
    row = 1
    if data['batch_dur']:
        ws6.cell(row=row, column=1, value='Длительность этапов партий').font = title_font
        row += 1
        write_header(ws6, row, ['Переход', 'Среднее время', 'Количество'])
        row += 1
        for i, d in enumerate(data['batch_dur']):
            write_row(ws6, row, [d['label'], d['avg'], d['count']], i % 2 == 1)
            row += 1
        row += 1

    if data['order_dur']:
        ws6.cell(row=row, column=1, value='Длительность этапов заказов').font = title_font
        row += 1
        write_header(ws6, row, ['Переход', 'Среднее время', 'Количество'])
        row += 1
        for i, d in enumerate(data['order_dur']):
            write_row(ws6, row, [d['label'], d['avg'], d['count']], i % 2 == 1)
            row += 1
    auto_width(ws6)

    # ---- Sheet 7: Отгруженные заказы ----
    ws7 = wb.create_sheet('Отгруженные заказы')
    ws7.cell(row=1, column=1, value=f'Отгруженные заказы за период ({period})').font = title_font
    ws7.cell(row=2, column=1, value=f'Всего: {len(data["shipped_orders"])}').font = green_font

    write_header(ws7, 4, ['Номер заказа', 'Приоритет', 'Создан', 'Позиции (товар)', 'Позиции (SKU)', 'Позиции (кол-во)'])
    row = 5
    for order in data['shipped_orders']:
        lines = order.lines.select_related('product').all()
        if lines:
            for j, line in enumerate(lines):
                vals = [
                    order.order_number if j == 0 else '',
                    order.get_priority_code_display() if j == 0 else '',
                    order.created_at.strftime('%d.%m.%Y %H:%M') if j == 0 else '',
                    line.product.product_name,
                    line.product.sku_code,
                    _qty_str(line.requested_qty),
                ]
                write_row(ws7, row, vals, (row - 5) % 2 == 1)
                row += 1
        else:
            write_row(ws7, row, [
                order.order_number, order.get_priority_code_display(),
                order.created_at.strftime('%d.%m.%Y %H:%M'), '—', '—', '—',
            ])
            row += 1
    auto_width(ws7)

    # ---- Sheet 8: Просроченные заказы ----
    ws8 = wb.create_sheet('Просроченные заказы')
    ws8.cell(row=1, column=1, value='Просроченные заказы').font = title_font
    ws8.cell(row=2, column=1, value=f'Всего: {len(data["overdue_orders"])}').font = red_font

    write_header(ws8, 4, ['Номер заказа', 'Приоритет', 'Текущий этап', 'Плановая отгрузка', 'Просрочка'])
    for i, o in enumerate(data['overdue_orders']):
        write_row(ws8, 5 + i, [
            o.order_number,
            o.get_priority_code_display(),
            o.get_current_stage_code_display(),
            o.planned_ship_date.strftime('%d.%m.%Y %H:%M') if o.planned_ship_date else '—',
            _delay_str(now, o.planned_ship_date),
        ], i % 2 == 1)

    # Детализация позиций под основной таблицей
    row = 5 + len(data['overdue_orders']) + 2
    for order in data['overdue_orders']:
        ws8.cell(row=row, column=1, value=f'Заказ {order.order_number}').font = Font(
            name='Arial', size=11, bold=True, color='1A1F36')
        row += 1
        write_header(ws8, row, ['Товар', 'SKU', 'Количество', '', ''])
        row += 1
        lines = order.lines.select_related('product').all()
        for j, line in enumerate(lines):
            write_row(ws8, row, [
                line.product.product_name, line.product.sku_code,
                _qty_str(line.requested_qty), '', '',
            ], j % 2 == 1)
            row += 1
        row += 1
    auto_width(ws8)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
